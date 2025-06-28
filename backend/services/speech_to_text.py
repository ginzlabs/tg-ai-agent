import os
import asyncio
import aiohttp
from typing import Dict, Any, Optional, Tuple
from dotenv import load_dotenv
import json
from utils.logger import logger
from docx import Document
import sys
from config import STT_SUMMARY_TEXTS

# Add path for summarization service
sys.path.append(os.path.abspath(os.path.dirname(__file__)))
from summarization import SummarizationService

# Load environment variables
load_dotenv()

# Singleton instance
_stt_instance = None

class SpeechToText:
    """
    A lean speech-to-text class using AssemblyAI API with direct HTTP calls.
    Supports both local audio files and URLs.

    Usage:
    >>> stt = SpeechToText()
    >>> transcript_id, result = await stt.transcribe_async("https://example.com/audio.mp3", t_id="record_uuid")  # URL with default 'nano' model
    >>> transcript_id, result = await stt.transcribe_async("local_audio.mp3", speaker_labels=True, t_id="record_uuid")  # Local file with speaker labels
    >>> transcript_id, result = await stt.transcribe_async("audio.mp3", model="best", t_id="record_uuid")  # Use more accurate 'best' model
    >>> transcript_id, result = await stt.transcribe_async("audio.mp3", language_detection=False, language_code="fr", t_id="record_uuid")  # Disable language detection and use specific language
    """

    def __init__(self):
        """Initialize the SpeechToText class with AssemblyAI credentials."""
        self._setup_credentials()
        self.base_url = "https://api.assemblyai.com/v2"
        self.default_model = "nano"  # Use nano model by default for faster processing
        self.webhook_url = os.getenv("BACKEND01_WEBHOOK_URL")
        self.webhook_secret = os.getenv("OUR_SECRET_TOKEN")
        
        if not self.webhook_url:
            logger.warning("WEBHOOK_URL not found in environment variables. Webhook functionality will be disabled.")
        if not self.webhook_secret:
            logger.warning("WEBHOOK_SECRET not found in environment variables. Webhook security will be disabled.")

    def _setup_credentials(self):
        """Set up AssemblyAI API key."""
        self.api_key = os.getenv("ASSEMBLYAI_API_KEY")
        if not self.api_key:
            raise ValueError("ASSEMBLYAI_API_KEY not found in environment variables.")
        self.headers = {
            "authorization": self.api_key,
            "content-type": "application/json"
        }

    def _is_url(self, audio_url: str) -> bool:
        """Check if the input is a URL."""
        return audio_url.startswith("http://") or audio_url.startswith("https://")

    async def _upload_file_async(self, file_path: str) -> str:
        """
        Upload a local file to AssemblyAI asynchronously.
        
        Args:
            file_path: Path to the local audio file
            
        Returns:
            str: The upload_url for the file
        """
        logger.info(f"Uploading local file asynchronously: {file_path}")
        
        with open(file_path, "rb") as f:
            file_data = f.read()
        
        async with aiohttp.ClientSession() as session:
            async with session.post(
                f"{self.base_url}/upload",
                headers={"authorization": self.api_key},
                data=file_data
            ) as response:
                if response.status != 200:
                    error_text = await response.text()
                    raise RuntimeError(f"Upload failed: {error_text}")
                
                response_json = await response.json()
                return response_json["upload_url"]

    async def get_transcript_by_id(self, transcript_id: str) -> Dict[str, Any]:
        """
        Get a transcript using its ID.
        
        Args:
            transcript_id: The AssemblyAI transcript ID
            
        Returns:
            Dict[str, Any]: The transcript data
        """
        logger.info(f"Getting transcript by ID: {transcript_id}")
        
        async with aiohttp.ClientSession() as session:
            async with session.get(
                f"{self.base_url}/transcript/{transcript_id}",
                headers=self.headers
            ) as response:
                if response.status != 200:
                    error_text = await response.text()
                    raise RuntimeError(f"Failed to get transcript {transcript_id}: {error_text}")
                
                return await response.json()

    async def transcribe_async(self, audio_url: str, speaker_labels: bool = False, model: Optional[str] = None, 
                              language_detection: bool = True, language_code: Optional[str] = None,
                              t_id: Optional[str] = None, chat_id: Optional[int] = None, 
                              db_thread_id: Optional[str] = None, message_id: Optional[int] = None, temp_msg_id: Optional[int] = None) -> Tuple[str, Dict[str, Any]]:
        """
        Asynchronously transcribe an audio file or URL using AssemblyAI with webhook support.

        Args:
            audio_url (str): Path to the audio file or URL (e.g., "https://example.com/audio.mp3").
            speaker_labels (bool): Whether to enable speaker diarization. Default is False.
            model (Optional[str]): Model to use for transcription. Options: "nano" (faster), "best" (more accurate). Default is "nano".
            language_detection (bool): Whether to enable automatic language detection. Default is True.
            language_code (Optional[str]): Specific language code to use (e.g., "en_us", "fr", "es"). If specified and language_detection is True, it will be used as a fallback.
            t_id (Optional[str]): Optional ID to associate with this transcription in the webhook URL.
            chat_id (Optional[int]): Optional chat ID to include in webhook URL parameters.
            db_thread_id (Optional[str]): Optional DB thread ID to include in webhook URL parameters.
            message_id (Optional[int]): Optional message ID to include in webhook URL parameters.
            temp_msg_id (Optional[int]): Optional temporary message ID to include in webhook URL parameters.
        Returns:
            Tuple[str, Dict[str, Any]]: The transcript ID and a partial result dictionary.

        Raises:
            Exception: If transcription request fails.
        """
        # If audio_url is a local file, upload it first
        if not self._is_url(audio_url):
            audio_url_to_use = await self._upload_file_async(audio_url)
        else:
            audio_url_to_use = audio_url
        
        # Prepare transcription request
        data = {
            "audio_url": audio_url_to_use,
            "speaker_labels": speaker_labels,
            "language_detection": language_detection,
            "speech_model": model or self.default_model,  # Use provided model or default to 'nano'
        }
        
        if language_code:
            data["language_code"] = language_code
        
        # Add webhook configuration if available
        if self.webhook_url:
            webhook_url = self.webhook_url
            # Build query parameters for the webhook URL
            query_params = []
            
            if t_id:
                query_params.append(f"t_id={t_id}")
                
            if chat_id is not None:
                query_params.append(f"chat_id={chat_id}")

            if db_thread_id is not None:
                query_params.append(f"db_thread_id={db_thread_id}")
                
            if message_id is not None:
                query_params.append(f"message_id={message_id}")

            if temp_msg_id is not None:
                query_params.append(f"temp_msg_id={temp_msg_id}")
                
            # Construct the full webhook URL with parameters
            if query_params:
                webhook_url += f"/webhooks/assemblyai?{'&'.join(query_params)}"
            else:
                webhook_url += "/webhooks/assemblyai"
            
            data["webhook_url"] = webhook_url
            
            if self.webhook_secret:
                data["webhook_auth_header_name"] = "X-Webhook-Secret"
                data["webhook_auth_header_value"] = self.webhook_secret
        
        # Start transcription
        async with aiohttp.ClientSession() as session:
            async with session.post(
                f"{self.base_url}/transcript",
                json=data,
                headers=self.headers
            ) as response:
                if response.status != 200:
                    error_text = await response.text()
                    raise RuntimeError(f"Transcription request failed: {error_text}")
                
                response_json = await response.json()
                transcript_id = response_json['id']
        
        # Return partial result immediately
        partial_result = {
            "transcript_id": transcript_id,
            "status": "queued",
            "metadata": {
                "audio_source": audio_url,
                "provider": "AssemblyAI",
                "diarization": speaker_labels,
                "model": model or self.default_model,
                "language_detection": language_detection,
                "language_code": language_code or "auto-detect",
                "webhook_enabled": bool(self.webhook_url)
            }
        }
        
        logger.info(f"Started transcription with ID {transcript_id} using webhook: {bool(self.webhook_url)}")
        return transcript_id, partial_result

    async def format_transcript_result(self, transcription_result: Dict[str, Any], speaker_labels: bool = False, 
                                     language_detection: bool = True, model: Optional[str] = None) -> Dict[str, Any]:
        """
        Format the transcript result into a standardized structure.
        
        Args:
            transcription_result: The raw transcription result from AssemblyAI.
            speaker_labels: Whether speaker labels were requested.
            language_detection: Whether language detection was enabled.
            model: The model used for transcription.
            
        Returns:
            Dict[str, Any]: Formatted transcript result.
        """
        # Format the result
        result = {
            "text": transcription_result.get("text", ""),
            "transcript_id": transcription_result.get("id"),
            "status": transcription_result.get("status")
        }
        
        # Include speakers data if speaker labels were requested and available
        # Check if the transcription result contains utterances and not empty and if the text is more than 1000 chars
        if "utterances" in transcription_result and transcription_result["utterances"] and len(result["text"]) > 1000:
            try:
                # Create an in-memory DOCX file from utterances
                from io import BytesIO
                from docx import Document
                
                # Function to convert JSON to the requested text format
                json_data = transcription_result["utterances"]
                
                # Generate transcript with speakers first
                transcript_with_speakers = ""
                for entry in json_data:
                    transcript_with_speakers += f"[Speaker {entry['speaker']}]:\n{entry['text']}\n\n"
                
                # Store transcript with speakers in result
                result["transcript_with_speakers"] = transcript_with_speakers
                logger.info(f"Created transcript with speakers for transcript {result['transcript_id']}")
                
                # Initialize summarization service with timeout
                summarization_service = SummarizationService(timeout=60)
                
                summary_result = None
                try:
                    logger.info(f"Generating summary for transcript {result['transcript_id']}")
                    # Use transcript with speakers for summarization
                    summary_result = await summarization_service.summarize(transcript_with_speakers)
                    # Include the summary in the result
                    result["summary"] = summary_result.get("summary", "")
                    logger.info(f"Summary generation completed in {summary_result.get('timing', {}).get('total_seconds', 0)} seconds")
                except Exception as e:
                    logger.error(f"Failed to generate summary: {str(e)}", exc_info=True)
            
                # Create a new Document
                doc = Document()
                
                # Add summary section if available
                if summary_result and summary_result.get("summary") and not summary_result.get("timing", {}).get("timed_out", False):
                    # Add title for summary
                    summary_title = doc.add_heading(STT_SUMMARY_TEXTS["summary"]["title"], level=1)
                    
                    # Add summary text
                    doc.add_paragraph(summary_result.get("summary", ""))
                    
                    # Add key points if available
                    if summary_result.get("key_points"):
                        doc.add_heading(STT_SUMMARY_TEXTS["summary"]["key_points_title"], level=2)
                        for point in summary_result.get("key_points", []):
                            bullet_item = doc.add_paragraph(style='List Bullet')
                            bullet_item.add_run(point)
                    
                    # Add transcript title
                    doc.add_heading(STT_SUMMARY_TEXTS["summary"]["full_transcript_title"], level=1)
                
                # Add utterances with speaker labels
                for entry in json_data:
                    # Add speaker paragraph
                    speaker_para = doc.add_paragraph()
                    speaker_para.add_run(STT_SUMMARY_TEXTS["speaker"]["label"].format(entry['speaker'])).bold = True
                    # Add text paragraph
                    doc.add_paragraph(entry['text'])
                    # Add empty paragraph for spacing
                    doc.add_paragraph()
                
                # Save to BytesIO object
                docx_buffer = BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                
                # Get bytes from BytesIO for storage
                result["transcription_docx"] = docx_buffer.getvalue()
                logger.info(f"Created in-memory DOCX file for transcript {result['transcript_id']} (size: {len(result['transcription_docx'])} bytes)")
                
            except Exception as e:
                logger.error(f"Failed to create DOCX from utterances: {str(e)}", exc_info=True)
                # Don't raise the exception, just log it and continue

        # Include detected language information if available
        if language_detection and "language_code" in transcription_result:
            result["detected_language"] = {
                "language_code": transcription_result["language_code"]
            }
            
            # Include confidence score if available
            if "language_confidence" in transcription_result:
                result["detected_language"]["confidence"] = transcription_result["language_confidence"]
        
        # Include the model used
        result["model_used"] = transcription_result.get("speech_model", model or self.default_model)
        
        return result

def get_stt_instance() -> SpeechToText:
    """Get or create the SpeechToText instance (singleton pattern)"""
    global _stt_instance
    if _stt_instance is None:
        _stt_instance = SpeechToText()
    return _stt_instance

async def process_speech_to_text(audio_url: str, speaker_labels: bool = False, model: Optional[str] = None,
                                language_detection: bool = True, language_code: Optional[str] = None,
                                t_id: Optional[str] = None, chat_id: Optional[int] = None,
                                db_thread_id: Optional[str] = None, message_id: Optional[int] = None, temp_msg_id: Optional[int] = None) -> Dict[str, Any]:
    """
    Process speech to text transcription for the given audio input.
    
    Args:
        audio_url: URL or path to the audio file
        speaker_labels: Whether to enable speaker diarization. Default is False.
        model: Model to use for transcription. Options: "nano" (faster), "best" (more accurate). Default is "nano".
        language_detection: Whether to enable automatic language detection. Default is True.
        language_code: Specific language code to use. Default is None (auto-detect or use AssemblyAI's default).
        t_id: Optional ID to associate with this transcription in the webhook URL.
        chat_id: Chat ID to include in webhook URL parameters.
        db_thread_id: DB thread ID to include in webhook URL parameters.
        message_id: Message ID to include in webhook URL parameters.
        temp_msg_id: Temporary message ID to include in webhook URL parameters.

    Returns:
        Dict containing the transcription result or initial response with transcript_id
    """
    stt = get_stt_instance()
    actual_model = model or stt.default_model
    
    logger.info(f"Starting speech-to-text processing for {audio_url} (speaker_labels={speaker_labels}, model={actual_model}, language_detection={language_detection})")
    
    try:
        transcript_id, partial_result = await stt.transcribe_async(
            audio_url, 
            speaker_labels=speaker_labels, 
            model=model,
            language_detection=language_detection,
            language_code=language_code,
            t_id=t_id,
            chat_id=chat_id,
            db_thread_id=db_thread_id,
            message_id=message_id,
            temp_msg_id=temp_msg_id
        )
        
        logger.info(f"Initiated speech-to-text processing for {audio_url} with transcript_id: {transcript_id}")
        return {
            "transcript_id": transcript_id,
            **partial_result
        }
    except Exception as e:
        logger.error(f"Speech-to-text processing failed: {str(e)}", exc_info=True)
        raise

async def get_transcript_result(transcript_id: str, speaker_labels: bool = False, 
                               language_detection: bool = True, model: Optional[str] = None) -> Dict[str, Any]:
    """
    Retrieve and format transcription result for a given transcript ID.
    
    Args:
        transcript_id: The AssemblyAI transcript ID.
        speaker_labels: Whether speaker labels were requested.
        language_detection: Whether language detection was enabled.
        model: The model used for transcription.
        
    Returns:
        Dict containing the formatted transcription result.
    """
    stt = get_stt_instance()
    
    try:
        transcription_result = await stt.get_transcript_by_id(transcript_id)
        formatted_result = await stt.format_transcript_result(
            transcription_result,
            speaker_labels=speaker_labels,
            language_detection=language_detection,
            model=model
        )

        # print(formatted_result["speakers"])
        
        return formatted_result
    except Exception as e:
        logger.error(f"Failed to get transcript result: {str(e)}", exc_info=True)
        raise 